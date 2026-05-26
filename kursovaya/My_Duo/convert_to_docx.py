import os
import sys
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

template_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\My_Duo\Титульный КР 2026.docx"
chapters_dir = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\My_Duo\chapters"
output_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\My_Duo\Шакшуев_КР_Отчет.docx"

if not os.path.exists(template_path):
    print(f"Error: Template file {template_path} not found!")
    sys.exit(1)

print("Opening template Word document...")
doc = docx.Document(template_path)

# --- ROBUST SEARCH-AND-REPLACE ON TITLE PAGE ---
print("Applying search-and-replace to customize Title Page details...")

replacements = {
    "Копаницкий Захар Александрович": "Шакшуев Роман Викторович",
    "Копаницкого Захара Александровича": "Шакшуева Романа Викторовича",
    "Копаницкий З.А.": "Шакшуев Р.В.",
    "Б9123-09.03.03пи": "Б9123-09.03.03пикд",
    "Разработка модуля мониторинга Telegram-бота для автоматического сбора и региональной сортировки новостей": "Создание Telegram-бота: разработка пользовательского интерфейса и инструментов для администрирования системы мониторинга",
    "Разработке модуля мониторинга Telegram-бота для автоматического сбора и региональной сортировки новостей": "Созданию Telegram-бота: разработка пользовательского интерфейса и инструментов для администрирования системы мониторинга"
}

def replace_text_in_runs(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Simplest way: if the whole old_text is in one run, or if we do it at the paragraph level
        # To avoid breaking run formatting, we can check if it exists in individual runs first
        replaced = False
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced = True
        
        # If it spanned multiple runs, replace at paragraph level and keep formatting on first run
        if not replaced and paragraph.runs:
            p_text = paragraph.text.replace(old_text, new_text)
            paragraph.text = ""
            run = paragraph.add_run(p_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

# Apply replacement to all paragraphs in the document (including headers/footers)
for p in doc.paragraphs:
    for old, new in replacements.items():
        replace_text_in_runs(p, old, new)

# Apply replacement to all table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in replacements.items():
                    replace_text_in_runs(p, old, new)

print("Title Page details successfully replaced!")


bookmark_id_counter = 0

def add_bookmark_to_paragraph(paragraph, name):
    global bookmark_id_counter
    b_id = bookmark_id_counter
    bookmark_id_counter += 1
    
    bmStart = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{b_id}" w:name="{name}"/>')
    bmEnd = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{b_id}"/>')
    
    paragraph._p.insert(0, bmStart)
    paragraph._p.append(bmEnd)

def add_hyperlink_to_bookmark(paragraph, text, anchor, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=(0, 51, 153)):
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} w:anchor="{anchor}">\n'
        f'  <w:r>\n'
        f'    <w:rPr>\n'
        f'      <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>\n'
        f'      <w:sz w:val="{int(size_pt * 2)}"/>\n'
        f'      <w:szCs w:val="{int(size_pt * 2)}"/>\n'
        f'      {"<w:b/>" if bold else ""}\n'
        f'      {"<w:i/>" if italic else ""}\n'
        f'      <w:color w:val="{"%02X%02X%02X" % color_rgb}"/>\n'
        f'      <w:u w:val="single"/>\n'\
        f'    </w:rPr>\n'
        f'    <w:t>{text}</w:t>\n'
        f'  </w:r>\n'
        f'</w:hyperlink>'
    )
    paragraph._p.append(hyperlink)
    return hyperlink

def set_paragraph_heading_style(p, level):
    if level == 1:
        style_name = 'Заголовок 11'
    elif level == 2:
        style_name = 'Заголовок 21'
    else:
        return
        
    doc_obj = p.part.document
    
    try:
        doc_obj.styles[style_name]
    except KeyError:
        try:
            new_style = doc_obj.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            new_style.base_style = doc_obj.styles['Normal']
            
            # Configure style properties
            if level == 1:
                new_style.font.name = 'Times New Roman'
                new_style.font.size = Pt(16)
                new_style.font.bold = True
                new_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_style.paragraph_format.line_spacing = 1.5
                new_style.paragraph_format.space_before = Pt(12)
                new_style.paragraph_format.space_after = Pt(12)
            elif level == 2:
                new_style.font.name = 'Times New Roman'
                new_style.font.size = Pt(14)
                new_style.font.bold = True
                new_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_style.paragraph_format.line_spacing = 1.5
                new_style.paragraph_format.space_before = Pt(12)
                new_style.paragraph_format.space_after = Pt(6)
        except Exception as e:
            print(f"Error creating style {style_name}: {e}")
            
    try:
        p.style = style_name
    except Exception as e:
        print(f"Error assigning style {style_name}: {e}")
        
    # Assign Outline Level for Navigation Pane (0-indexed)
    try:
        pPr = p._p.get_or_add_pPr()
        outlineLvl = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
        pPr.append(outlineLvl)
    except Exception as e:
        print(f"Warning setting outline level: {e}")

# Helper function to set formatting on runs
def format_run(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=(0, 0, 0)):
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

# Helper function to style paragraphs
def style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=1.25, space_before=0, space_after=6):
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    if first_indent_cm > 0:
        p.paragraph_format.first_line_indent = Inches(first_indent_cm * 0.3937)
    else:
        p.paragraph_format.first_line_indent = Inches(0.0)

# Helper to shade code blocks
def shade_paragraph(p, fill_color="F4F4F4"):
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    pPr.append(shd)

# Helper to set borders on tables manually
def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def parse_markdown_table(lines, start_idx):
    headers = []
    rows = []
    
    header_line = lines[start_idx]
    headers = [c.strip() for c in header_line.split('|')[1:-1]]
    
    idx = start_idx + 2
    
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_line = lines[idx].strip()
        if not re.match(r'^\|[\s:-|-]*\|$', row_line):
            cells = [c.strip() for c in row_line.split('|')[1:-1]]
            rows.append(cells)
        idx += 1
        
    return headers, rows, idx

def add_table_to_word(doc, headers, rows):
    num_cols = len(headers)
    num_rows = len(rows) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    set_table_borders(table)
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        style_paragraph(p, align="left", line_spacing=1.15, first_indent_cm=0, space_after=2)
        run = p.add_run(title)
        format_run(run, size_pt=11, bold=True)
        
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF1F5"/>'))
        
    # Set data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_value in enumerate(row_data):
            val = cell_value if c_idx < len(row_data) else ""
            val_clean = val.replace("**", "").replace("*", "")
            
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            style_paragraph(p, align="left", line_spacing=1.15, first_indent_cm=0, space_after=2)
            run = p.add_run(val_clean)
            
            # Bold highlights
            is_bold = "Разрабатываемый" in val or "Наш" in val or "Мгновенная" in val or "Высокая" in val or "Автоматическая" in val
            format_run(run, size_pt=10, bold=is_bold)
            
            # Shading for our solution
            if r_idx == 4 or "Разрабатываемый" in row_data[0] or "Наш" in row_data[0]:
                tcPr = row_cells[c_idx]._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F9F2"/>'))

    empty_p = doc.add_paragraph()
    style_paragraph(empty_p, first_indent_cm=0, space_after=12)

def append_markdown_file(doc, file_path):
    file_name = os.path.basename(file_path)
    print(f"Reading and parsing {file_name}...")
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split('\n')
    
    in_code_block = False
    code_text = []
    
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        line_strip = line.strip()
        
        # 0. Mermaid block boundary
        if line_strip.startswith('```mermaid'):
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith('```'):
                idx += 1
            idx += 1
            
            # We embed the dynamically generated admin flow diagram
            image_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\My_Duo\admin_flowchart.png"
                
            if os.path.exists(image_path):
                p_img = doc.add_paragraph()
                style_paragraph(p_img, align="center", first_indent_cm=0, space_before=12, space_after=6)
                r_img = p_img.add_run()
                r_img.add_picture(image_path, width=Inches(6.0))
            else:
                print(f"Warning: Image {image_path} not found!")
            continue
            
        # 1. Code block boundary
        if line_strip.startswith('```'):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                style_paragraph(p, align="left", line_spacing=1.0, first_indent_cm=0, space_before=4, space_after=6)
                shade_paragraph(p)
                
                # Add borders
                pPr = p._p.get_or_add_pPr()
                pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>'))
                
                code_content = "\n".join(code_text)
                run = p.add_run(code_content)
                format_run(run, font_name="Consolas", size_pt=9.5, color_rgb=(60, 60, 60))
                code_text = []
            else:
                in_code_block = True
            idx += 1
            continue
            
        if in_code_block:
            code_text.append(line)
            idx += 1
            continue
            
        if line_strip == "---" or line_strip == "":
            idx += 1
            continue
            
        # 3. Tables parsing
        if line_strip.startswith('|'):
            headers, rows, next_idx = parse_markdown_table(lines, idx)
            add_table_to_word(doc, headers, rows)
            idx = next_idx
            continue
            
        # 4. Headings
        if line_strip.startswith('# '):
            title = line_strip[2:]
            doc.add_page_break()
            
            p = doc.add_paragraph()
            set_paragraph_heading_style(p, 1)
            style_paragraph(p, align="center", first_indent_cm=0, space_before=12, space_after=12)
            run = p.add_run(title.upper())
            format_run(run, size_pt=16, bold=True)
            
            # Insert Bookmark
            heading_lower = title.lower()
            anchor_name = None
            if "введение" in heading_lower: anchor_name = "vvedenie"
            elif "проектирование" in heading_lower: anchor_name = "glava1"
            elif "реализация" in heading_lower: anchor_name = "glava2"
            elif "заключение" in heading_lower: anchor_name = "zakluchenie"
            elif "литературы" in heading_lower: anchor_name = "literatura"
            elif "приложение" in heading_lower: anchor_name = "prilozhenie"
                
            if anchor_name:
                add_bookmark_to_paragraph(p, anchor_name)
                
            idx += 1
            continue
            
        if line_strip.startswith('## '):
            title = line_strip[3:]
            p = doc.add_paragraph()
            set_paragraph_heading_style(p, 2)
            style_paragraph(p, align="left", first_indent_cm=0, space_before=12, space_after=6)
            run = p.add_run(title)
            format_run(run, size_pt=14, bold=True)
            
            # Insert Bookmark
            heading_lower = title.lower()
            anchor_name = None
            if "1.1" in heading_lower: anchor_name = "sec1_1"
            elif "1.2" in heading_lower: anchor_name = "sec1_2"
            elif "1.3" in heading_lower: anchor_name = "sec1_3"
            elif "1.4" in heading_lower: anchor_name = "sec1_4"
            elif "1.5" in heading_lower: anchor_name = "sec1_5"
            elif "1.6" in heading_lower: anchor_name = "sec1_6"
            elif "2.1" in heading_lower: anchor_name = "sec2_1"
            elif "2.2" in heading_lower: anchor_name = "sec2_2"
            elif "2.3" in heading_lower: anchor_name = "sec2_3"
            elif "2.4" in heading_lower: anchor_name = "sec2_4"
            elif "2.5" in heading_lower: anchor_name = "sec2_5"
            elif "2.6" in heading_lower: anchor_name = "sec2_6"
            
            if anchor_name:
                add_bookmark_to_paragraph(p, anchor_name)
                
            idx += 1
            continue
            
        if line_strip.startswith('### '):
            title = line_strip[4:]
            p = doc.add_paragraph()
            style_paragraph(p, align="left", first_indent_cm=1.25, space_before=6, space_after=4)
            run = p.add_run(title)
            format_run(run, size_pt=14, bold=True, italic=False)
            idx += 1
            continue

        if line_strip.startswith('#### '):
            title = line_strip[5:]
            title = title.replace("`", "")
            p = doc.add_paragraph()
            style_paragraph(p, align="left", first_indent_cm=1.25, space_before=6, space_after=4)
            run = p.add_run(title)
            format_run(run, size_pt=14, bold=True, italic=False)
            idx += 1
            continue
            
        # 5. Lists (Bullet/Numbered)
        match_num = re.match(r'^(\d+)\.\s+(.*)', line_strip)
        if match_num:
            num = match_num.group(1)
            list_text = match_num.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=0, space_after=4)
            
            r_num = p.add_run(f"{num}. ")
            format_run(r_num, size_pt=14, bold=True)
            
            parse_inline_text(p, list_text)
            idx += 1
            continue
            
        if line_strip.startswith('* ') or line_strip.startswith('- ') or line_strip.startswith('• '):
            list_text = line_strip[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=0, space_after=4)
            
            r_bullet = p.add_run("•  ")
            format_run(r_bullet, size_pt=14, bold=True)
            
            parse_inline_text(p, list_text)
            idx += 1
            continue
            
        # Markdown image parsing: ![alt](url)
        match_img = re.match(r'^!\[(.*?)\]\((.*?)\)', line_strip)
        if match_img:
            alt_text = match_img.group(1)
            img_url = match_img.group(2)
            
            img_path = img_url.replace("file:///", "").replace("/", "\\")
            if not (img_path.lower().startswith("c:\\") or os.path.exists(img_path)):
                img_path = os.path.join(os.path.dirname(output_path), os.path.basename(img_path))
                
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                style_paragraph(p_img, align="center", first_indent_cm=0, space_before=12, space_after=6)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(6.0))
                print(f"Embedded markdown image: {img_path}")
            else:
                print(f"Warning: Markdown image file not found at {img_path}!")
            idx += 1
            continue
            
        # Figure caption (CENTER, 12pt, normal, no italic, space before 4, space after 12)
        if line_strip.startswith('Рисунок'):
            p = doc.add_paragraph()
            style_paragraph(p, align="center", line_spacing=1.5, first_indent_cm=0, space_before=4, space_after=12)
            run = p.add_run(line_strip)
            format_run(run, size_pt=12, bold=False, italic=False)
            idx += 1
            continue
            
        # Table caption (CENTER, 12pt, normal, 1.25 indent, space before 12, space after 6)
        if line_strip.startswith('Таблица'):
            p = doc.add_paragraph()
            style_paragraph(p, align="center", line_spacing=1.5, first_indent_cm=1.25, space_before=12, space_after=6)
            run = p.add_run(line_strip)
            format_run(run, size_pt=12, bold=False, italic=False)
            idx += 1
            continue
            
        # 6. Standard Paragraphs
        p = doc.add_paragraph()
        style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=1.25, space_after=6)
        parse_inline_text(p, line_strip)
        idx += 1

def parse_inline_text(p, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = p.add_run(inner)
            format_run(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            inner = part[1:-1]
            run = p.add_run(inner)
            format_run(run, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            run = p.add_run(inner)
            format_run(run, font_name="Consolas", size_pt=None, color_rgb=(0, 0, 0)) # Inherited size (14.0 pt)
        else:
            run = p.add_run(part)
            format_run(run)


# --- CONSTRUCT CHAPTERS SEQUENCE ---
chapter_files = [
    os.path.join(chapters_dir, "01_introduction_scoping.md"),
    os.path.join(chapters_dir, "02_architecture_ui.md"),
    os.path.join(chapters_dir, "03_realtime_ui.md"),
    os.path.join(chapters_dir, "04_administration.md")
]

# Insert page break after Title Page
doc.add_page_break()

# Insert СОДЕРЖАНИЕ Header
p = doc.add_paragraph()
style_paragraph(p, align="center", first_indent_cm=0, space_before=12, space_after=12)
run = p.add_run("СОДЕРЖАНИЕ")
format_run(run, size_pt=16, bold=True)

# TOC Configuration
toc_items = [
    ("ВВЕДЕНИЕ", "vvedenie", 3),
    ("1 ПРОЕКТИРОВАНИЕ СИСТЕМЫ", "glava1", 5),
    ("  1.1 Анализ предметной области и UX/UI паттернов", "sec1_1", 5),
    ("  1.2 Обоснование выбора методов и средств разработки", "sec1_2", 8),
    ("  1.3 Требования к аппаратному и программному окружению системы", "sec1_3", 10),
    ("  1.4 Проектирование карты экранов (User Flow) и структуры команд", "sec1_4", 12),
    ("  1.5 Логика безопасности и разграничения прав доступа", "sec1_5", 14),
    ("  1.6 Спецификация хранения данных пользователя и подписок", "sec1_6", 16),
    ("2 РЕАЛИЗАЦИЯ ПРОЕКТА", "glava2", 18),
    ("  2.1 Реализация асинхронного ядра бота", "sec2_1", 18),
    ("  2.2 Разработка модуля клавиатур и интерактивной навигации", "sec2_2", 21),
    ("  2.3 Обработка Callback-запросов и динамическое обновление интерфейса", "sec2_3", 24),
    ("  2.4 Панель администратора и инструменты управления каналами", "sec2_4", 27),
    ("  2.5 Реализация системы представления новостей и генерации дайджестов", "sec2_5", 30),
    ("  2.6 Тестирование юзабилити и сценариев взаимодействия", "sec2_6", 33),
    ("ЗАКЛЮЧЕНИЕ", "zakluchenie", 36),
    ("СПИСОК ЛИТЕРАТУРЫ", "literatura", 37),
    ("ПРИЛОЖЕНИЕ. Руководство по развертыванию и настройке бота", "prilozhenie", 39)
]

for title, anchor, page_num in toc_items:
    p_toc = doc.add_paragraph()
    style_paragraph(p_toc, align="left", line_spacing=1.15, first_indent_cm=0, space_after=3)
    p_toc.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT, docx.enum.text.WD_TAB_LEADER.DOTS)
    
    title_bold = title.strip() in ["ВВЕДЕНИЕ", "1 ПРОЕКТИРОВАНИЕ СИСТЕМЫ", "2 РЕАЛИЗАЦИЯ ПРОЕКТА", "ЗАКЛЮЧЕНИЕ", "СПИСОК ЛИТЕРАТУРЫ"] or "ПРИЛОЖЕНИЕ" in title
    
    # Clickable title with dot leader
    add_hyperlink_to_bookmark(p_toc, f"{title}\t", anchor, size_pt=12, bold=title_bold)
    # Page number
    add_hyperlink_to_bookmark(p_toc, str(page_num), anchor, size_pt=12, bold=title_bold)

# Append chapters
for file_path in chapter_files:
    if os.path.exists(file_path):
        append_markdown_file(doc, file_path)
    else:
        print(f"Warning: Chapter {file_path} not found!")

# Save final document
print(f"Saving compiled Word document to {output_path}...")
doc.save(output_path)
print("SUCCESS: Shakshuev's coursework report compiled into premium docx!")
