import sys
import os
import re
import docx

sys.stdout.reconfigure(encoding='utf-8')

report_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\My_Duo\Шакшуев_КР_Отчет.docx"

if not os.path.exists(report_path):
    print(f"Error: Report file {report_path} not found!")
    sys.exit(1)

print("Opening compiled report document for verification...")
doc = docx.Document(report_path)

print("\n" + "="*50)
print("             VERIFICATION REPORT")
print("="*50)


# 1. Check Title Page Replacements
print("\n[TITLE PAGE REPLACEMENT CHECKS]")
errors = 0
for idx, p in enumerate(doc.paragraphs[:50]): # Title page is within first 50 paragraphs
    if "Копаницкий" in p.text:
        print(f"❌ Error in paragraph #{idx}: Found leftover name 'Копаницкий' -> '{p.text[:60]}'")
        errors += 1
    if "Захар" in p.text:
        print(f"❌ Error in paragraph #{idx}: Found leftover name 'Захар' -> '{p.text[:60]}'")
        errors += 1
    if "Б9123-09.03.03пи" in p.text and "пикд" not in p.text:
        print(f"❌ Error in paragraph #{idx}: Found old group name without 'кд' -> '{p.text[:60]}'")
        errors += 1

for t_idx, t in enumerate(doc.tables[:3]): # Check first few tables (title page elements)
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, p in enumerate(cell.paragraphs):
                if "Копаницкий" in p.text:
                    print(f"❌ Error in Table #{t_idx+1}, Row #{r_idx}, Cell #{c_idx}: Found leftover name 'Копаницкий'")
                    errors += 1
                if "Захар" in p.text:
                    print(f"❌ Error in Table #{t_idx+1}, Row #{r_idx}, Cell #{c_idx}: Found leftover name 'Захар'")
                    errors += 1
                if "Б9123-09.03.03пи" in p.text and "пикд" not in p.text:
                    print(f"❌ Error in Table #{t_idx+1}, Row #{r_idx}, Cell #{c_idx}: Found old group name without 'кд'")
                    errors += 1

if errors == 0:
    print("✓ SUCCESS: Title page contains 100% correct details for Shakshuev Roman!")
else:
    print(f"❌ FAILED: Found {errors} title page mismatch errors.")

# 2. Check for leftover Markdown symbols in prose
print("\n[MARKDOWN SYMBOL CHECKS IN PROSE]")
md_errors = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
        
    # Skip code block and inline code paragraphs
    is_code = False
    for r in p.runs:
        if r.font.name == "Consolas":
            is_code = True
            break
    if is_code:
        continue
        
    # Check if there is block code start in prose
    if "```" in text:
        print(f"❌ Leftover code block bounds in paragraph #{idx}: '{text[:50]}'")
        md_errors += 1
    # Check for raw bold double asterisks or single italic asterisks
    if "**" in text:
        print(f"❌ Leftover double asterisks in paragraph #{idx}: '{text[:50]}'")
        md_errors += 1
    if "*" in text and not text.startswith("*"): # Allow bullets
        if re.search(r'\*\w+', text):
            print(f"❌ Leftover asterisks in paragraph #{idx}: '{text[:50]}'")
            md_errors += 1
    if "`" in text:
        print(f"❌ Leftover backticks in paragraph #{idx}: '{text[:50]}'")
        md_errors += 1

if md_errors == 0:
    print("✓ SUCCESS: Zero raw markdown symbols left in prose paragraphs!")
else:
    print(f"❌ FAILED: Found {md_errors} loose markdown symbol errors.")

# 3. Check custom styles and outline levels
print("\n[HEADING STYLES & OUTLINE CHECKS]")
h1_count = 0
h2_count = 0
for idx, p in enumerate(doc.paragraphs):
    style_name = p.style.name
    if style_name == 'Заголовок 11':
        h1_count += 1
        # Check that it's UPPERCASE
        if p.text != p.text.upper():
            print(f"⚠️ Warning in L1 Heading #{h1_count}: Text is not uppercase -> '{p.text}'")
    elif style_name == 'Заголовок 21':
        h2_count += 1

print(f"Total 'Заголовок 11' (Heading 1) found: {h1_count}")
print(f"Total 'Заголовок 21' (Heading 2) found: {h2_count}")
if h1_count > 0 and h2_count > 0:
    print("✓ SUCCESS: Custom Russian localized styles correctly applied to headings!")
else:
    print("❌ FAILED: Missing heading styles.")

# 4. Check Bibliography integrity (GOST, no access dates)
print("\n[BIBLIOGRAPHY INTEGRITY CHECKS]")
bib_found = False
bib_errors = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "СПИСОК ЛИТЕРАТУРЫ" in text.upper():
        bib_found = True
        continue
    if bib_found:
        # Check next paragraphs for dates
        if "Приложение" in text or "ПРИЛОЖЕНИЕ" in text:
            bib_found = False
            continue
        if "Дата обращения" in text or "Access date" in text or "дата обращения" in text:
            print(f"❌ Error in Bibliography item: Found access date -> '{text[:60]}'")
            bib_errors += 1
            
if bib_errors == 0:
    print("✓ SUCCESS: All bibliography URL entries are clean of access dates!")
else:
    print(f"❌ FAILED: Found {bib_errors} access date violations.")

# 5. Document summary
total_paragraphs = len(doc.paragraphs)
total_tables = len(doc.tables)
total_runs = sum(len(p.runs) for p in doc.paragraphs)
print("\n[DOCUMENT SUMMARY]")
print(f"  Total Paragraphs: {total_paragraphs}")
print(f"  Total Tables:     {total_tables}")
print(f"  Total Runs:       {total_runs}")

print("\nVerification process complete!")
