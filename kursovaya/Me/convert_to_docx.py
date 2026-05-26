import os
import sys
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

template_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\Титульный КР 2026.docx"
chapters_dir = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\chapters"
output_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\Копаницкий_КР_Отчет.docx"

if not os.path.exists(template_path):
    print(f"Error: Template file {template_path} not found!")
    sys.exit(1)

print("Opening template Word document...")
doc = docx.Document(template_path)

bookmark_id_counter = 0

def add_bookmark_to_paragraph(paragraph, name):
    global bookmark_id_counter
    b_id = bookmark_id_counter
    bookmark_id_counter += 1
    
    bmStart = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{b_id}" w:name="{name}"/>')
    bmEnd = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{b_id}"/>')
    
    paragraph._p.insert(0, bmStart)
    paragraph._p.append(bmEnd)

def add_hyperlink_to_bookmark(paragraph, text, anchor, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=(0, 51, 153)):
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} w:anchor="{anchor}">\n'
        f'  <w:r>\n'
        f'    <w:rPr>\n'
        f'      <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>\n'
        f'      <w:sz w:val="{int(size_pt * 2)}"/>\n'
        f'      <w:szCs w:val="{int(size_pt * 2)}"/>\n'
        f'      {"<w:b/>" if bold else ""}\n'
        f'      {"<w:i/>" if italic else ""}\n'
        f'      <w:color w:val="{"%02X%02X%02X" % color_rgb}"/>\n'
        f'      <w:u w:val="single"/>\n'\
        f'    </w:rPr>\n'
        f'    <w:t>{text}</w:t>\n'
        f'  </w:r>\n'
        f'</w:hyperlink>'
    )
    paragraph._p.append(hyperlink)
    return hyperlink

def set_paragraph_heading_style(p, level):
    style_name = f'Heading {level}'
    
    # 1. Ensure style exists in document styles
    try:
        p.part.document.styles[style_name]
    except KeyError:
        try:
            # Create style dynamically based on 'Normal'
            new_style = p.part.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            new_style.base_style = p.part.document.styles['Normal']
        except Exception:
            pass
            
    # 2. Assign style
    try:
        p.style = style_name
    except KeyError:
        try:
            p.style = f'Заголовок {level}'
        except KeyError:
            pass
            
    # 3. Explicitly set OpenXML Outline Level for Navigation Pane (0-indexed)
    try:
        pPr = p._p.get_or_add_pPr()
        outlineLvl = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
        pPr.append(outlineLvl)
    except Exception as e:
        print(f"Warning setting outline level: {e}")

# Helper function to set formatting on runs
def format_run(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=(0, 0, 0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

# Helper function to style paragraphs
def style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=1.25, space_before=0, space_after=6):
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    if first_indent_cm > 0:
        # 1 cm = 0.3937 inches
        p.paragraph_format.first_line_indent = Inches(first_indent_cm * 0.3937)

# Helper to shade code blocks
def shade_paragraph(p, fill_color="F4F4F4"):
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    pPr.append(shd)

# Helper to set borders on tables manually
def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def parse_markdown_table(lines, start_idx):
    headers = []
    rows = []
    
    # Parse headers
    header_line = lines[start_idx]
    headers = [c.strip() for c in header_line.split('|')[1:-1]]
    
    # Skip separator line (e.g. |---|---|)
    idx = start_idx + 2
    
    # Parse data rows
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_line = lines[idx].strip()
        # Ensure it's not a separator line
        if not re.match(r'^\|[\s:-|-]*\|$', row_line):
            cells = [c.strip() for c in row_line.split('|')[1:-1]]
            rows.append(cells)
        idx += 1
        
    return headers, rows, idx

def add_table_to_word(doc, headers, rows):
    num_cols = len(headers)
    num_rows = len(rows) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    set_table_borders(table)
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        style_paragraph(p, align="left", line_spacing=1.15, first_indent_cm=0, space_after=2)
        run = p.add_run(title)
        format_run(run, size_pt=11, bold=True)
        
        # Shade header cell gray
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF1F5"/>'))
        
    # Set data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_value in enumerate(row_data):
            # Safe boundary check
            val = cell_value if c_idx < len(row_data) else ""
            
            # Clean bold text inside markdown table cells
            val_clean = val.replace("**", "").replace("*", "")
            
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            style_paragraph(p, align="left", line_spacing=1.15, first_indent_cm=0, space_after=2)
            run = p.add_run(val_clean)
            
            # Bold important lines
            is_bold = "Разрабатываемое" in val or "Полностью" in val or "Мгновенная" in val or "Автоматическая" in val or "Высокая" in val
            format_run(run, size_pt=10, bold=is_bold)
            
            # Highlight our solution row
            if r_idx == 3 or "Разрабатываемое" in row_data[0]:
                tcPr = row_cells[c_idx]._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F9F2"/>'))

    # Add space after table
    empty_p = doc.add_paragraph()
    style_paragraph(empty_p, first_indent_cm=0, space_after=12)

def append_markdown_file(doc, file_path):
    file_name = os.path.basename(file_path)
    print(f"Reading and parsing {file_name}...")
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split('\n')
    
    in_code_block = False
    code_text = []
    
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        line_strip = line.strip()
        
        # 0. Mermaid block boundary
        if line_strip.startswith('```mermaid'):
            # Skip the lines of the mermaid block until the closing ```
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith('```'):
                idx += 1
            idx += 1 # skip closing ```
            
            # Select the appropriate generated image based on the chapter file name
            if "02_architecture_database" in file_name:
                image_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\architecture_diagram.png"
            elif "04_algorithms_optimization" in file_name:
                image_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\algorithm_diagram.png"
            else:
                image_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\architecture_diagram.png"
                
            if os.path.exists(image_path):
                p_img = doc.add_paragraph()
                style_paragraph(p_img, align="center", first_indent_cm=0, space_before=12, space_after=6)
                r_img = p_img.add_run()
                r_img.add_picture(image_path, width=Inches(6.0))
            else:
                print(f"Warning: Image {image_path} not found!")
            continue
            
        # 1. Code block boundary
        if line_strip.startswith('```'):
            if in_code_block:
                # Close code block
                in_code_block = False
                p = doc.add_paragraph()
                style_paragraph(p, align="left", line_spacing=1.0, first_indent_cm=0, space_before=4, space_after=6)
                shade_paragraph(p)
                
                # Add thin border to code block paragraph
                pPr = p._p.get_or_add_pPr()
                pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>'))
                
                code_content = "\n".join(code_text)
                run = p.add_run(code_content)
                format_run(run, font_name="Consolas", size_pt=9.5, color_rgb=(60, 60, 60))
                code_text = []
            else:
                in_code_block = True
            idx += 1
            continue
            
        if in_code_block:
            code_text.append(line)
            idx += 1
            continue
            
        # 2. Skip dividers or separators
        if line_strip == "---" or line_strip == "":
            idx += 1
            continue
            
        # 3. Tables parsing
        if line_strip.startswith('|'):
            # Parse table
            headers, rows, next_idx = parse_markdown_table(lines, idx)
            add_table_to_word(doc, headers, rows)
            idx = next_idx
            continue
            
        # 4. Headings
        if line_strip.startswith('# '):
            title = line_strip[2:]
            
            # Start main chapters on a new page (e.g. Введение, Проектирование, Реализация, Заключение, Литература)
            doc.add_page_break()
            
            p = doc.add_paragraph()
            set_paragraph_heading_style(p, 1)
            # Centered for main headers, no indent, space before/after
            style_paragraph(p, align="center", first_indent_cm=0, space_before=12, space_after=12)
            run = p.add_run(title.upper())
            format_run(run, size_pt=16, bold=True)
            
            # Insert Bookmark
            heading_lower = title.lower()
            anchor_name = None
            if "введение" in heading_lower:
                anchor_name = "vvedenie"
            elif "проектирование" in heading_lower:
                anchor_name = "glava1"
            elif "реализация" in heading_lower:
                anchor_name = "glava2"
            elif "заключение" in heading_lower:
                anchor_name = "zakluchenie"
            elif "литературы" in heading_lower or "источников" in heading_lower:
                anchor_name = "literatura"
            elif "приложение" in heading_lower:
                anchor_name = "prilozhenie"
                
            if anchor_name:
                add_bookmark_to_paragraph(p, anchor_name)
                
            idx += 1
            continue
            
        if line_strip.startswith('## '):
            title = line_strip[3:]
            p = doc.add_paragraph()
            set_paragraph_heading_style(p, 2)
            style_paragraph(p, align="left", first_indent_cm=0, space_before=12, space_after=6)
            run = p.add_run(title)
            format_run(run, size_pt=14, bold=True)
            
            # Insert Bookmark
            heading_lower = title.lower()
            anchor_name = None
            if "1.1" in heading_lower: anchor_name = "sec1_1"
            elif "1.2" in heading_lower: anchor_name = "sec1_2"
            elif "1.3" in heading_lower: anchor_name = "sec1_3"
            elif "1.4" in heading_lower: anchor_name = "sec1_4"
            elif "1.5" in heading_lower: anchor_name = "sec1_5"
            elif "2.1" in heading_lower: anchor_name = "sec2_1"
            elif "2.2" in heading_lower: anchor_name = "sec2_2"
            elif "2.3" in heading_lower: anchor_name = "sec2_3"
            elif "2.4" in heading_lower: anchor_name = "sec2_4"
            elif "2.5" in heading_lower: anchor_name = "sec2_5"
            elif "2.6" in heading_lower: anchor_name = "sec2_6"
            elif "введение" in heading_lower: anchor_name = "vvedenie"
            elif "заключение" in heading_lower: anchor_name = "zakluchenie"
            elif "список литературы" in heading_lower: anchor_name = "literatura"
            
            if anchor_name:
                add_bookmark_to_paragraph(p, anchor_name)
                
            idx += 1
            continue
            
        if line_strip.startswith('### '):
            title = line_strip[4:]
            p = doc.add_paragraph()
            # Level 3 subheadings are styled as standard bold paragraph separators (not Headings)
            style_paragraph(p, align="left", first_indent_cm=1.25, space_before=6, space_after=4)
            run = p.add_run(title)
            format_run(run, size_pt=14, bold=True, italic=False)
            idx += 1
            continue

        if line_strip.startswith('#### '):
            title = line_strip[5:]
            title = title.replace("`", "") # Remove backticks in heading text
            p = doc.add_paragraph()
            # Level 4 subheadings are styled as standard bold paragraph separators (not Headings)
            style_paragraph(p, align="left", first_indent_cm=1.25, space_before=6, space_after=4)
            run = p.add_run(title)
            format_run(run, size_pt=14, bold=True, italic=False)
            idx += 1
            continue
            
        # 5. Lists (Bullet/Numbered)
        # Handle numbered lists like "1. ", "2. "
        match_num = re.match(r'^(\d+)\.\s+(.*)', line_strip)
        if match_num:
            num = match_num.group(1)
            list_text = match_num.group(2)
            p = doc.add_paragraph()
            # Left indent 0.75 cm, no first line indent
            p.paragraph_format.left_indent = Inches(0.3)
            style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=0, space_after=4)
            
            # Highlight number
            r_num = p.add_run(f"{num}. ")
            format_run(r_num, size_pt=14, bold=True)
            
            # Parse bold/italic elements inside text
            parse_inline_text(p, list_text)
            idx += 1
            continue
            
        # Handle bullet lists like "* ", "- ", "• "
        if line_strip.startswith('* ') or line_strip.startswith('- ') or line_strip.startswith('• '):
            list_text = line_strip[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=0, space_after=4)
            
            # Add bullet symbol
            r_bullet = p.add_run("•  ")
            format_run(r_bullet, size_pt=14, bold=True)
            
            parse_inline_text(p, list_text)
            idx += 1
            continue
            
        # Check if it's a figure caption (e.g. starts with "*Рисунок" or "Рисунок")
        if (line_strip.startswith('*Рисунок') and line_strip.endswith('*')) or line_strip.startswith('Рисунок'):
            caption_text = line_strip
            if caption_text.startswith('*') and caption_text.endswith('*'):
                caption_text = caption_text[1:-1]
            
            p = doc.add_paragraph()
            style_paragraph(p, align="center", first_indent_cm=0, space_before=4, space_after=12)
            run = p.add_run(caption_text)
            format_run(run, size_pt=12, italic=True)
            idx += 1
            continue
            
        # 6. Standard Paragraphs
        p = doc.add_paragraph()
        style_paragraph(p, align="justify", line_spacing=1.5, first_indent_cm=1.25, space_after=6)
        parse_inline_text(p, line_strip)
        idx += 1

def parse_inline_text(p, text):
    # Parse bold text "**", italic text "*", and inline code "`"
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = p.add_run(inner)
            format_run(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            inner = part[1:-1]
            run = p.add_run(inner)
            format_run(run, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            run = p.add_run(inner)
            # Use Consolas for clean code look without raw backticks
            format_run(run, font_name="Consolas", size_pt=11, color_rgb=(60, 60, 60))
        else:
            run = p.add_run(part)
            format_run(run)

# 2. Append Chapters
# Order: Introduction/Scoping, Architecture/Database, Realtime/Parser, Algorithms/Optimization
chapter_files = [
    os.path.join(chapters_dir, "01_introduction_scoping.md"),
    os.path.join(chapters_dir, "02_architecture_database.md"),
    os.path.join(chapters_dir, "03_realtime_parser.md"),
    os.path.join(chapters_dir, "04_algorithms_optimization.md")
]

# Before appending chapters, let's insert a page break to separate from template's Title Page!
doc.add_page_break()

# Insert Содержание placeholder
p = doc.add_paragraph()
style_paragraph(p, align="center", first_indent_cm=0, space_before=12, space_after=12)
run = p.add_run("СОДЕРЖАНИЕ")
format_run(run, size_pt=16, bold=True)

# Add hyperlinked TOC matching markdown structure
toc_items = [
    ("ВВЕДЕНИЕ", "vvedenie", 3),
    ("1 ПРОЕКТИРОВАНИЕ СИСТЕМЫ", "glava1", 5),
    ("  1.1 Анализ предметной области и существующих решений", "sec1_1", 5),
    ("  1.2 Обоснование выбора методов и средств разработки", "sec1_2", 7),
    ("  1.3 Требования к аппаратному и программному окружению системы", "sec1_3", 9),
    ("  1.4 Общая архитектура системы сбора данных", "sec1_4", 11),
    ("  1.5 Спецификация данных и проектирование базы данных", "sec1_5", 13),
    ("2 РЕАЛИЗАЦИЯ ПРОЕКТА", "glava2", 16),
    ("  2.1 Реализация клиента-сборщика реального времени", "sec2_1", 16),
    ("  2.2 Модуль первичной обработки новостей и парсинга", "sec2_2", 19),
    ("  2.3 Алгоритм очистки текста от форматирования", "sec2_3", 22),
    ("  2.4 Алгоритм региональной сортировки и классификации", "sec2_4", 24),
    ("  2.5 Алгоритмы дедупликации и технической фильтрации", "sec2_5", 27),
    ("  2.6 Оптимизация производительности при высокой нагрузке", "sec2_6", 30),
    ("ЗАКЛЮЧЕНИЕ", "zakluchenie", 33),
    ("СПИСОК ЛИТЕРАТУРЫ", "literatura", 34),
    ("ПРИЛОЖЕНИЕ А. Руководство по развертыванию и запуску", "prilozhenie", 36)
]

for title, anchor, page_num in toc_items:
    p_toc = doc.add_paragraph()
    style_paragraph(p_toc, align="left", line_spacing=1.15, first_indent_cm=0, space_after=3)
    p_toc.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT, docx.enum.text.WD_TAB_LEADER.DOTS)
    
    title_bold = title.strip() == "ВВЕДЕНИЕ" or title.strip() == "1 ПРОЕКТИРОВАНИЕ СИСТЕМЫ" or title.strip() == "2 РЕАЛИЗАЦИЯ ПРОЕКТА" or title.strip() == "ЗАКЛЮЧЕНИЕ" or title.strip() == "СПИСОК ЛИТЕРАТУРЫ" or "ПРИЛОЖЕНИЕ А" in title
    
    # Add a clickable hyperlinked title to its bookmark
    add_hyperlink_to_bookmark(p_toc, f"{title}\t", anchor, size_pt=12, bold=title_bold)
    
    # Page number hyperlink
    add_hyperlink_to_bookmark(p_toc, str(page_num), anchor, size_pt=12, bold=title_bold)

# Append each chapter file
for file_path in chapter_files:
    if os.path.exists(file_path):
        append_markdown_file(doc, file_path)
    else:
        print(f"Warning: File {file_path} not found!")

# Save docx
print(f"Saving compiled Word document to {output_path}...")
doc.save(output_path)
print("SUCCESS: Coursework report compiled into Word document!")
