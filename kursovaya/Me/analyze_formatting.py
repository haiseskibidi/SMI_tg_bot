import os
import sys
import docx
from docx.shared import Pt, Inches

sys.stdout.reconfigure(encoding='utf-8')

user_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\ОТЧЁТ.docx"
our_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\Копаницкий_КР_Отчет.docx"

if not os.path.exists(user_doc_path):
    print(f"Error: User's finalized document not found at {user_doc_path}")
    sys.exit(1)

if not os.path.exists(our_doc_path):
    print(f"Error: Our compiled document not found at {our_doc_path}")
    sys.exit(1)

def analyze_document(path):
    doc = docx.Document(path)
    info = {
        "paragraphs_count": len(doc.paragraphs),
        "headings": {},
        "normal_text": [],
        "list_items": [],
        "tables": []
    }
    
    # 1. Analyze page margins of first section
    section = doc.sections[0]
    info["margins"] = {
        "top": section.top_margin.inches if section.top_margin else None,
        "bottom": section.bottom_margin.inches if section.bottom_margin else None,
        "left": section.left_margin.inches if section.left_margin else None,
        "right": section.right_margin.inches if section.right_margin else None,
    }
    
    # 2. Analyze paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        style_name = p.style.name
        
        # Determine runs formatting
        runs_info = []
        for r in p.runs:
            runs_info.append({
                "text": r.text,
                "font_name": r.font.name,
                "font_size": r.font.size.pt if r.font.size else None,
                "bold": r.bold,
                "italic": r.italic,
                "color": r.font.color.rgb if r.font.color else None
            })
            
        p_format = {
            "alignment": str(p.alignment) if p.alignment else "None",
            "line_spacing": p.paragraph_format.line_spacing,
            "first_line_indent": p.paragraph_format.first_line_indent.inches if p.paragraph_format.first_line_indent else None,
            "left_indent": p.paragraph_format.left_indent.inches if p.paragraph_format.left_indent else None,
            "space_before": p.paragraph_format.space_before.pt if p.paragraph_format.space_before else None,
            "space_after": p.paragraph_format.space_after.pt if p.paragraph_format.space_after else None,
            "style": style_name,
            "runs": runs_info,
            "text_preview": text[:60]
        }
        
        # Classify by style name or level
        if "Heading" in style_name or "Заголовок" in style_name:
            info["headings"][style_name] = info["headings"].get(style_name, [])
            info["headings"][style_name].append(p_format)
        elif "List" in style_name or p_format["left_indent"] is not None or text.startswith("•") or (text and text[0].isdigit() and ". " in text[:5]):
            info["list_items"].append(p_format)
        else:
            info["normal_text"].append(p_format)
            
    # 3. Analyze tables
    for t_idx, t in enumerate(doc.tables):
        rows_info = []
        for r_idx, row in enumerate(t.rows):
            cells_info = []
            for c_idx, cell in enumerate(row.cells):
                p = cell.paragraphs[0] if cell.paragraphs else None
                runs_data = []
                if p:
                    for r in p.runs:
                        runs_data.append({
                            "font_name": r.font.name,
                            "font_size": r.font.size.pt if r.font.size else None,
                            "bold": r.bold,
                            "italic": r.italic
                        })
                cells_info.append({
                    "text": cell.text.strip(),
                    "runs": runs_data
                })
            rows_info.append(cells_info)
        info["tables"].append({
            "rows_count": len(t.rows),
            "cols_count": len(t.columns),
            "rows": rows_info
        })
        
    return info

print("Analyzing User's finalized document (ОТЧЁТ.docx)...")
user_info = analyze_document(user_doc_path)

print("Analyzing Our compiled document (Копаницкий_КР_Отчет.docx)...")
our_info = analyze_document(our_doc_path)

# --- COMPARISON REPORT ---
print("\n" + "="*50)
print("             FORMATTING COMPARISON REPORT")
print("="*50)

# 1. Margins Compare
print("\n[PAGE MARGINS]")
print(f"         {'User (Gold)':<20} | {'Our Compiled':<20}")
for margin in ["left", "right", "top", "bottom"]:
    user_val = user_info["margins"][margin]
    our_val = our_info["margins"][margin]
    user_cm = f"{user_val * 2.54:.2f} cm" if user_val else "None"
    our_cm = f"{our_val * 2.54:.2f} cm" if our_val else "None"
    status = "✓ MATCH" if abs((user_val or 0) - (our_val or 0)) < 0.05 else "❌ MISMATCH"
    print(f"{margin:<8}: {user_cm:<20} | {our_cm:<20} -> {status}")

# 2. Paragraphs Count
print(f"\nParagraphs count: User={user_info['paragraphs_count']}, Ours={our_info['paragraphs_count']}")

# 3. Normal Text formatting
print("\n[NORMAL PARAGRAPH FORMATTING]")
if user_info["normal_text"]:
    up = user_info["normal_text"][0]
    op = our_info["normal_text"][0] if our_info["normal_text"] else None
    
    # Extract run properties
    urun = up["runs"][0] if up["runs"] else {}
    orun = op["runs"][0] if op and op["runs"] else {}
    
    print(f"Font Name:      User={urun.get('font_name')} | Ours={orun.get('font_name')}")
    print(f"Font Size:      User={urun.get('font_size')} pt | Ours={orun.get('font_size')} pt")
    print(f"Alignment:      User={up['alignment']} | Ours={op['alignment'] if op else 'None'}")
    print(f"Line Spacing:   User={up['line_spacing']} | Ours={op['line_spacing'] if op else 'None'}")
    
    u_indent = f"{up['first_line_indent'] * 2.54:.2f} cm" if up['first_line_indent'] else "None"
    o_indent = f"{op['first_line_indent'] * 2.54:.2f} cm" if op and op['first_line_indent'] else "None"
    print(f"First-line Ind: User={u_indent} | Ours={o_indent}")
    print(f"Space Before:   User={up['space_before']} pt | Ours={op['space_before'] if op else 'None'} pt")
    print(f"Space After:    User={up['space_after']} pt | Ours={op['space_after'] if op else 'None'} pt")

# 4. Heading formatting compare
print("\n[HEADING STYLES]")
all_headings = set(user_info["headings"].keys()) | set(our_info["headings"].keys())
for h_style in sorted(all_headings):
    u_list = user_info["headings"].get(h_style, [])
    o_list = our_info["headings"].get(h_style, [])
    
    print(f"\nStyle: {h_style}")
    if u_list:
        up = u_list[0]
        urun = up["runs"][0] if up["runs"] else {}
        print(f"  User (Gold):   Size={urun.get('font_size')} pt | Bold={urun.get('bold')} | Italic={urun.get('italic')} | Align={up['alignment']} | Before={up['space_before']} pt | After={up['space_after']} pt")
    else:
        print("  User (Gold):   Not used!")
        
    if o_list:
        op = o_list[0]
        orun = op["runs"][0] if op["runs"] else {}
        print(f"  Our Compiled:  Size={orun.get('font_size')} pt | Bold={orun.get('bold')} | Italic={orun.get('italic')} | Align={op['alignment']} | Before={op['space_before']} pt | After={op['space_after']} pt")
    else:
        print("  Our Compiled:  Not used!")

# 5. Check for direct run modifications (e.g. did user remove italic, or bold somewhere?)
print("\n[ITALIC & BOLD COMPARISON IN RUNS]")
# Let's count how many italics or bolds exist in user's doc compared to ours
u_italic_runs = 0
u_bold_runs = 0
for p in user_info["normal_text"] + user_info["list_items"]:
    for r in p["runs"]:
        if r["italic"]: u_italic_runs += 1
        if r["bold"]: u_bold_runs += 1
        
o_italic_runs = 0
o_bold_runs = 0
for p in our_info["normal_text"] + our_info["list_items"]:
    for r in p["runs"]:
        if r["italic"]: o_italic_runs += 1
        if r["bold"]: o_bold_runs += 1

print(f"Total Italic runs: User={u_italic_runs} | Ours={o_italic_runs}")
print(f"Total Bold runs:   User={u_bold_runs} | Ours={o_bold_runs}")

# 6. Compare first 10 headings text to see if there are structural differences
print("\n[HEADING TEXTS & ANCHOR NAMES]")
user_headings_text = []
for h_style in sorted(user_info["headings"].keys()):
    for p in user_info["headings"][h_style]:
        user_headings_text.append((h_style, p["text_preview"]))

print("User Headings (Top 10):")
for style, text in user_headings_text[:10]:
    print(f"  [{style}]: {text}")
