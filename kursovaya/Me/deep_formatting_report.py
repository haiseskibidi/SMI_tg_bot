import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

user_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\ОТЧЁТ.docx"

doc = docx.Document(user_doc_path)

print("="*60)
print("DEEP FORMATTING ANALYSIS OF USER'S GOLD STANDARD (ОТЧЁТ.docx)")
print("="*60)

# 1. Analyze first 20 normal paragraphs to find the actual font and spacing
normal_count = 0
print("\n[NORMAL PARAGRAPHS DETAILED FORMATTING]")
for p in doc.paragraphs:
    # Skip empty paragraphs and headings
    if not p.text.strip():
        continue
    style_name = p.style.name
    if "Заголовок" in style_name or "Heading" in style_name:
        continue
        
    normal_count += 1
    if normal_count > 10:
        break
        
    # Get paragraph formatting
    pf = p.paragraph_format
    line_spacing = pf.line_spacing
    alignment = p.alignment
    first_line_indent = pf.first_line_indent.cm if pf.first_line_indent else 0.0
    space_before = pf.space_before.pt if pf.space_before else 0.0
    space_after = pf.space_after.pt if pf.space_after else 0.0
    
    # Extract run details
    run_details = []
    for r in p.runs[:3]: # look at first few runs
        font_name = r.font.name
        font_size = r.font.size.pt if r.font.size else "Default"
        bold = r.bold
        italic = r.italic
        run_details.append(f"Run('{r.text[:15]}', font={font_name}, size={font_size}, B={bold}, I={italic})")
        
    print(f"Paragraph {normal_count} (Preview: '{p.text[:40]}...'):")
    print(f"  Style: {style_name} | Align: {alignment} | LineSpacing: {line_spacing} | FirstLineIndent: {first_line_indent:.2f} cm")
    print(f"  Space Before: {space_before} pt | Space After: {space_after} pt")
    print(f"  Runs: {', '.join(run_details)}")

# 2. Analyze list items in user document
list_count = 0
print("\n[LIST PARAGRAPHS DETAILED FORMATTING]")
for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    pf = p.paragraph_format
    if pf.left_indent is not None or p.style.name.startswith("List") or text.startswith("•") or text.startswith("-"):
        list_count += 1
        if list_count > 5:
            break
        print(f"List Item {list_count} (Preview: '{text[:50]}...'):")
        print(f"  Style: {p.style.name} | LeftIndent: {pf.left_indent.cm if pf.left_indent else 0:.2f} cm | FirstLineIndent: {pf.first_line_indent.cm if pf.first_line_indent else 0:.2f} cm")
        for r in p.runs[:2]:
            print(f"    Run: '{r.text[:15]}' | font={r.font.name} | size={r.font.size.pt if r.font.size else 'Default'} | B={r.bold} | I={r.italic}")

# 3. Analyze code blocks
print("\n[CODE BLOCKS FORMATTING]")
code_block_count = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    # Let's detect code blocks by font name Consolas or Courier
    is_code = False
    for r in p.runs:
        if r.font.name in ["Consolas", "Courier New"]:
            is_code = True
            break
    if is_code or p.style.name == "Normal" and p.paragraph_format.left_indent and not text.startswith("•") and not text[0].isdigit():
        code_block_count += 1
        if code_block_count > 5:
            break
        pf = p.paragraph_format
        print(f"Code Paragraph {code_block_count} (Preview: '{text[:50]}...'):")
        print(f"  Style: {p.style.name} | Align: {p.alignment} | LeftIndent: {pf.left_indent.cm if pf.left_indent else 0:.2f} cm")
        for r in p.runs[:2]:
            print(f"    Run: '{r.text[:15]}' | font={r.font.name} | size={r.font.size.pt if r.font.size else 'Default'} | B={r.bold} | I={r.italic}")

# 4. Analyze table cells
print("\n[TABLES FORMATTING]")
for t_idx, t in enumerate(doc.tables[:2]):
    print(f"Table {t_idx+1}: Rows={len(t.rows)}, Cols={len(t.columns)}")
    for r_idx in range(min(len(t.rows), 3)):
        row = t.rows[r_idx]
        cells_preview = []
        for cell in row.cells[:3]:
            cells_preview.append(cell.text.strip().replace("\n", " ")[:30])
        print(f"  Row {r_idx+1}: {cells_preview}")
        # Look at formatting in cell 0
        if row.cells:
            cell = row.cells[0]
            if cell.paragraphs:
                p = cell.paragraphs[0]
                pf = p.paragraph_format
                print(f"    Cell 1 Paragraph Align: {p.alignment} | LineSpacing: {pf.line_spacing}")
                for r in p.runs[:2]:
                    print(f"      Run: font={r.font.name} | size={r.font.size.pt if r.font.size else 'Default'} | B={r.bold} | I={r.italic}")
