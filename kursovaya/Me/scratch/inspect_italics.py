import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

user_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\ОТЧЁТ.docx"
doc = docx.Document(user_doc_path)

print("ALL ITALIC RUNS IN USER'S DOC:")
print("-" * 60)
italic_count = 0
for idx, p in enumerate(doc.paragraphs):
    for r_idx, r in enumerate(p.runs):
        if r.italic:
            italic_count += 1
            print(f"Italic Run #{italic_count} in Paragraph #{idx} (style: {p.style.name}):")
            print(f"  Parent text preview: '{p.text.strip()[:100]}...'")
            print(f"  Italic run text: '{r.text}'")
            print(f"  Font: {r.font.name} | Size: {r.font.size.pt if r.font.size else None}")
            print()

# Check tables for italics too
for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, p in enumerate(cell.paragraphs):
                for r in p.runs:
                    if r.italic:
                        italic_count += 1
                        print(f"Italic Run #{italic_count} in Table #{t_idx+1}, Row #{r_idx}, Cell #{c_idx}:")
                        print(f"  Italic run text: '{r.text}'")
                        print(f"  Font: {r.font.name} | Size: {r.font.size.pt if r.font.size else None}")
                        print()

print(f"Total Italic runs found: {italic_count}")
