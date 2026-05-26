import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

user_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\ОТЧЁТ.docx"
doc = docx.Document(user_doc_path)

import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

user_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\ОТЧЁТ.docx"
doc = docx.Document(user_doc_path)

print("BLOCK-LEVEL CODE BLOCKS IN USER DOC:")
print("-" * 50)
block_count = 0
for idx, p in enumerate(doc.paragraphs):
    # Detect block-level code paragraphs by checking if they contain multiple lines
    # or have Consolas runs and are left-aligned or shaded.
    text = p.text.strip()
    if not text:
        continue
        
    has_consolas = False
    all_runs_consolas = True if p.runs else False
    for r in p.runs:
        if r.font.name == "Consolas":
            has_consolas = True
        else:
            all_runs_consolas = False
            
    # Typically, block code is entirely Consolas, or has a smaller font size (e.g. 9.5)
    is_block_code = has_consolas and (all_runs_consolas or "\n" in p.text or (p.runs and p.runs[0].font.size and p.runs[0].font.size.pt < 11))
    
    if is_block_code:
        block_count += 1
        if block_count > 10:
            print("... and more block code paragraphs")
            break
        print(f"\nBlock Code Paragraph #{idx}:")
        print(f"  Text (first 100 chars): {p.text.strip().replace('\n', ' ')[:100]}...")
        pf = p.paragraph_format
        print(f"  Style: {p.style.name} | Align: {p.alignment} | LineSpacing: {pf.line_spacing}")
        print(f"  FirstIndent: {pf.first_line_indent.cm if pf.first_line_indent else 0:.2f} cm | LeftIndent: {pf.left_indent.cm if pf.left_indent else 0:.2f} cm")
        print(f"  SpaceBefore: {pf.space_before.pt if pf.space_before else 0:.2f} pt | SpaceAfter: {pf.space_after.pt if pf.space_after else 0:.2f} pt")
        print(f"  Runs:")
        for r_idx, r in enumerate(p.runs[:3]):
            print(f"    Run {r_idx}: '{r.text[:40].replace('\n', ' ')}' | font={r.font.name} | size={r.font.size.pt if r.font.size else None} | B={r.bold} | I={r.italic}")

print(f"\nTotal Block Code paragraphs found: {block_count}")

