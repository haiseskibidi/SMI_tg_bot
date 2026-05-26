import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

user_doc_path = r"C:\Users\Q1\Documents\SMI_tg_bot\kursovaya\Me\ОТЧЁТ.docx"
doc = docx.Document(user_doc_path)

def print_p_info(idx, p, comment=""):
    pf = p.paragraph_format
    line_spacing = pf.line_spacing
    alignment = p.alignment
    first_line_indent = pf.first_line_indent.cm if pf.first_line_indent else 0.0
    left_indent = pf.left_indent.cm if pf.left_indent else 0.0
    space_before = pf.space_before.pt if pf.space_before else 0.0
    space_after = pf.space_after.pt if pf.space_after else 0.0
    
    runs_info = []
    for r in p.runs:
        runs_info.append({
            "text": r.text,
            "font": r.font.name,
            "size": r.font.size.pt if r.font.size else None,
            "bold": r.bold,
            "italic": r.italic
        })
        
    print(f"\n[{comment}] Paragraph #{idx} [Style: {p.style.name}]:")
    print(f"  Text: {p.text.strip()[:100]}...")
    print(f"  Align: {alignment} | LineSpacing: {line_spacing} | FirstLineIndent: {first_line_indent:.2f} cm | LeftIndent: {left_indent:.2f} cm")
    print(f"  SpaceBefore: {space_before} pt | SpaceAfter: {space_after} pt")
    print(f"  Runs count: {len(p.runs)}")
    for r_idx, r in enumerate(runs_info[:4]):
        print(f"    Run {r_idx}: '{r['text'][:40]}' | font={r['font']} | size={r['size']} | B={r['bold']} | I={r['italic']}")

print("FINDING STRUCTURAL ELEMENTS IN THE DOCUMENT")
print("=" * 60)

# Track counts of found items to avoid dumping too much
bullet_count = 0
code_count = 0
figure_count = 0
bib_found = False
bib_count = 0

for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
        
    # Check for bullet points
    if (text.startswith("•") or text.startswith("-") or p.style.name.startswith("List") or p.paragraph_format.left_indent) and bullet_count < 4:
        # Avoid title page entries which have left_indent or normal centered style
        if "Прикладная" not in text and "Владивосток" not in text and "2026" not in text and not text.startswith("1.") and not text.startswith("2."):
            bullet_count += 1
            print_p_info(idx, p, f"BULLET ITEM {bullet_count}")
            
    # Check for code blocks
    is_code = False
    for r in p.runs:
        if r.font.name in ["Consolas", "Courier New"]:
            is_code = True
            break
    if is_code and code_count < 3:
        code_count += 1
        print_p_info(idx, p, f"CODE BLOCK ITEM {code_count}")
        
    # Check for figure captions
    if ("рисунок" in text.lower() or "рис. " in text.lower()) and "рисунке" not in text.lower() and figure_count < 3:
        figure_count += 1
        print_p_info(idx, p, f"FIGURE CAPTION {figure_count}")
        
    # Check for bibliography
    if "СПИСОК ЛИТЕРАТУРЫ" in text.upper():
        bib_found = True
        print_p_info(idx, p, "BIBLIOGRAPHY HEADER")
        continue
        
    if bib_found:
        bib_count += 1
        if bib_count < 10:
            print_p_info(idx, p, f"BIBLIOGRAPHY ITEM {bib_count}")
        else:
            bib_found = False # Stop bibliography dumping

# Analyze the first table cell details
print("\n" + "=" * 60)
print("TABLES ANALYSIS")
print("=" * 60)
for t_idx, t in enumerate(doc.tables):
    print(f"\nTable #{t_idx+1}: Rows={len(t.rows)}, Cols={len(t.columns)}")
    # Look at header cell and body cell
    if len(t.rows) > 0 and len(t.columns) > 0:
        cell_0 = t.rows[0].cells[0]
        print(f"  Header Cell [0,0] Text: {cell_0.text.strip()}")
        if cell_0.paragraphs:
            p = cell_0.paragraphs[0]
            print_p_info(f"T{t_idx+1} [0,0]", p, "TABLE HEADER CELL")
            
    if len(t.rows) > 1 and len(t.columns) > 0:
        cell_1 = t.rows[1].cells[0]
        print(f"  Body Cell [1,0] Text: {cell_1.text.strip()}")
        if cell_1.paragraphs:
            p = cell_1.paragraphs[0]
            print_p_info(f"T{t_idx+1} [1,0]", p, "TABLE BODY CELL")
